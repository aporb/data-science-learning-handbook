#!/usr/bin/env python3
"""
Template Processor
==================

Advanced template processing engine for compliance documents with multi-format output,
digital signature support, and automated quality assurance.

Key Features:
- Multi-format output (HTML, PDF, DOCX, Markdown)
- Digital signature integration
- Automated quality assurance
- Content sanitization
- Template preprocessing and optimization
- Batch processing capabilities

Classification: UNCLASSIFIED//FOR OFFICIAL USE ONLY
Version: 1.0
Author: Security Compliance Team
Date: 2025-07-28
"""

import os
import json
import logging
import hashlib
import base64
from pathlib import Path
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, asdict
import concurrent.futures
from io import BytesIO

# Third-party imports with fallbacks
try:
    import pdfkit
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    pdfkit = None

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

try:
    import markdown
    from markdown.extensions import toc, tables, codehilite
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown = None

# Signature and encryption support
try:
    from cryptography.hazmat.primitives import hashes, serialization
    from cryptography.hazmat.primitives.asymmetric import rsa, padding
    from cryptography.hazmat.primitives.serialization import pkcs12
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False

from .compliance_template_engine import TemplateType, ClassificationLevel, ComplianceMetadata
from .template_validator import TemplateValidator, ValidationResult

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class OutputFormat:
    """Supported output formats"""
    HTML = "html"
    PDF = "pdf" 
    DOCX = "docx"
    MARKDOWN = "md"
    JSON = "json"


@dataclass
class ProcessingOptions:
    """Template processing options"""
    output_format: str = OutputFormat.HTML
    include_toc: bool = True
    include_signatures: bool = False
    sanitize_content: bool = True
    validate_output: bool = True
    compress_output: bool = False
    watermark_text: Optional[str] = None
    custom_css: Optional[str] = None
    page_numbering: bool = True
    classification_headers: bool = True
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return asdict(self)


@dataclass 
class ProcessingResult:
    """Template processing result"""
    success: bool
    output_path: Optional[str]
    output_format: str
    file_size: int
    processing_time: float
    content_hash: str
    validation_result: Optional[ValidationResult]
    errors: List[str]
    warnings: List[str]
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        result = asdict(self)
        if self.validation_result:
            result['validation_result'] = self.validation_result.to_dict()
        return result


class TemplateProcessor:
    """
    Advanced Template Processor for Compliance Documents
    
    Provides multi-format output, digital signatures, and quality assurance.
    """
    
    def __init__(self, 
                 output_path: Path,
                 config: Optional[Dict[str, Any]] = None):
        """
        Initialize Template Processor
        
        Args:
            output_path: Path for processed documents
            config: Configuration dictionary
        """
        self.output_path = Path(output_path)
        self.config = config or {}
        
        # Ensure output directory exists
        self.output_path.mkdir(parents=True, exist_ok=True)
        
        # Initialize validator
        self.validator = TemplateValidator()
        
        # Initialize format processors
        self._init_format_processors()
        
        # Digital signature setup
        self.signature_key = None
        self.signature_cert = None
        self._init_digital_signatures()
        
        logger.info(f"Template Processor initialized")
        logger.info(f"Output path: {self.output_path}")
        logger.info(f"Available formats: {self.get_available_formats()}")
    
    def _init_format_processors(self):
        """Initialize format-specific processors"""
        self.format_processors = {
            OutputFormat.HTML: self._process_html,
            OutputFormat.JSON: self._process_json,
        }
        
        if PDF_AVAILABLE:
            self.format_processors[OutputFormat.PDF] = self._process_pdf
        
        if DOCX_AVAILABLE:
            self.format_processors[OutputFormat.DOCX] = self._process_docx
        
        if MARKDOWN_AVAILABLE:
            self.format_processors[OutputFormat.MARKDOWN] = self._process_markdown
    
    def _init_digital_signatures(self):
        """Initialize digital signature capabilities"""
        if not CRYPTO_AVAILABLE:
            logger.warning("Cryptography not available, digital signatures disabled")
            return
        
        # Check for signature configuration
        sig_config = self.config.get('digital_signatures', {})
        key_path = sig_config.get('private_key_path')
        cert_path = sig_config.get('certificate_path')
        
        if key_path and cert_path:
            try:
                # Load private key
                with open(key_path, 'rb') as f:
                    self.signature_key = serialization.load_pem_private_key(
                        f.read(), 
                        password=None
                    )
                
                # Load certificate
                with open(cert_path, 'rb') as f:
                    self.signature_cert = serialization.load_pem_x509_certificate(f.read())
                
                logger.info("Digital signature capabilities enabled")
            except Exception as e:
                logger.warning(f"Could not load signature keys: {e}")
    
    def get_available_formats(self) -> List[str]:
        """Get list of available output formats"""
        return list(self.format_processors.keys())
    
    def sanitize_content(self, content: str) -> str:
        """
        Sanitize content for security
        
        Args:
            content: Content to sanitize
            
        Returns:
            Sanitized content
        """
        # Remove potentially dangerous script tags
        import re
        
        # Remove script tags
        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove javascript: URLs
        content = re.sub(r'javascript:', 'blocked:', content, flags=re.IGNORECASE)
        
        # Remove on* event handlers
        content = re.sub(r'\s+on\w+\s*=\s*["\'][^"\']*["\']', '', content, flags=re.IGNORECASE)
        
        # Remove style with javascript
        content = re.sub(r'style\s*=\s*["\'][^"\']*javascript[^"\']*["\']', '', content, flags=re.IGNORECASE)
        
        return content
    
    def add_classification_headers(self, 
                                 content: str, 
                                 classification: ClassificationLevel) -> str:
        """
        Add classification headers and footers
        
        Args:
            content: Document content
            classification: Classification level
            
        Returns:
            Content with classification markings
        """
        classification_text = f"CLASSIFICATION: {classification.value}"
        
        # Add header if HTML
        if '<body>' in content:
            header = f"""
            <div style="text-align: center; font-weight: bold; color: red; 
                       border: 2px solid red; padding: 10px; margin: 10px 0;">
                {classification_text}
            </div>
            """
            content = content.replace('<body>', f'<body>{header}')
            
            # Add footer
            footer = f"""
            <div style="text-align: center; font-weight: bold; color: red; 
                       border: 2px solid red; padding: 10px; margin: 10px 0;">
                {classification_text}
            </div>
            """
            content = content.replace('</body>', f'{footer}</body>')
        
        return content
    
    def add_watermark(self, content: str, watermark_text: str) -> str:
        """
        Add watermark to content
        
        Args:
            content: Document content
            watermark_text: Watermark text
            
        Returns:
            Content with watermark
        """
        if '<body>' in content:
            watermark_style = """
            <style>
            .watermark {
                position: fixed;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%) rotate(-45deg);
                font-size: 6em;
                color: rgba(200, 200, 200, 0.3);
                z-index: -1;
                pointer-events: none;
                user-select: none;
            }
            </style>
            """
            
            watermark_div = f"""
            <div class="watermark">{watermark_text}</div>
            """
            
            content = content.replace('</head>', f'{watermark_style}</head>')
            content = content.replace('<body>', f'<body>{watermark_div}')
        
        return content
    
    def _process_html(self, 
                     content: str,
                     metadata: ComplianceMetadata,
                     options: ProcessingOptions) -> Tuple[str, str]:
        """
        Process content to HTML format
        
        Args:
            content: Template content
            metadata: Document metadata
            options: Processing options
            
        Returns:
            Tuple of (processed_content, file_extension)
        """
        processed_content = content
        
        # Add custom CSS if provided
        if options.custom_css:
            css_style = f"<style>{options.custom_css}</style>"
            if '</head>' in processed_content:
                processed_content = processed_content.replace('</head>', f'{css_style}</head>')
        
        # Add classification headers
        if options.classification_headers:
            processed_content = self.add_classification_headers(
                processed_content, metadata.classification
            )
        
        # Add watermark
        if options.watermark_text:
            processed_content = self.add_watermark(processed_content, options.watermark_text)
        
        # Sanitize content if requested
        if options.sanitize_content:
            processed_content = self.sanitize_content(processed_content)
        
        return processed_content, "html"
    
    def _process_pdf(self,
                    content: str,
                    metadata: ComplianceMetadata, 
                    options: ProcessingOptions) -> Tuple[bytes, str]:
        """
        Process content to PDF format
        
        Args:
            content: HTML content
            metadata: Document metadata
            options: Processing options
            
        Returns:
            Tuple of (pdf_bytes, file_extension)
        """
        if not PDF_AVAILABLE:
            raise ValueError("PDF processing not available - pdfkit not installed")
        
        # Configure PDF options
        pdf_options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in', 
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': None
        }
        
        # Add page numbering if requested
        if options.page_numbering:
            pdf_options.update({
                'header-right': f'{metadata.system_name} - {metadata.classification.value}',
                'header-font-size': '9',
                'footer-center': '[page] of [topage]',
                'footer-font-size': '9'
            })
        
        try:
            # Generate PDF
            pdf_bytes = pdfkit.from_string(content, False, options=pdf_options)
            return pdf_bytes, "pdf"
        except Exception as e:
            raise ValueError(f"PDF generation failed: {str(e)}")
    
    def _process_docx(self,
                     content: str,
                     metadata: ComplianceMetadata,
                     options: ProcessingOptions) -> Tuple[bytes, str]:
        """
        Process content to DOCX format
        
        Args:
            content: HTML/text content
            metadata: Document metadata
            options: Processing options
            
        Returns:
            Tuple of (docx_bytes, file_extension)
        """
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing not available - python-docx not installed")
        
        # Create document
        doc = Document()
        
        # Add classification header
        if options.classification_headers:
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(f"CLASSIFICATION: {metadata.classification.value}")
            header_run.bold = True
            header_para.alignment = 1  # Center alignment
        
        # Add title
        title = doc.add_heading(f'{metadata.system_name} - {metadata.template_type.value.replace("_", " ").title()}', 0)
        
        # Add metadata table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        metadata_items = [
            ('System Name', metadata.system_name),
            ('System ID', metadata.system_id),
            ('Version', metadata.version),
            ('Created Date', metadata.created_date.strftime('%Y-%m-%d')),
            ('Classification', metadata.classification.value),
            ('Organization', metadata.organization)
        ]
        
        for key, value in metadata_items:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        # Simple content processing (HTML to text)
        import re
        text_content = re.sub(r'<[^>]+>', '', content)  # Strip HTML tags
        text_content = re.sub(r'\s+', ' ', text_content)  # Normalize whitespace
        
        # Add main content
        doc.add_paragraph(text_content)
        
        # Save to bytes
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()
        docx_buffer.close()
        
        return docx_bytes, "docx"
    
    def _process_markdown(self,
                         content: str,
                         metadata: ComplianceMetadata,
                         options: ProcessingOptions) -> Tuple[str, str]:
        """
        Process content to Markdown format
        
        Args:
            content: HTML content
            metadata: Document metadata
            options: Processing options
            
        Returns:
            Tuple of (markdown_content, file_extension)
        """
        if not MARKDOWN_AVAILABLE:
            raise ValueError("Markdown processing not available - markdown not installed")
        
        # Simple HTML to Markdown conversion
        import re
        
        # Convert HTML headers
        content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1\n', content, flags=re.IGNORECASE)
        content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1\n', content, flags=re.IGNORECASE)
        content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1\n', content, flags=re.IGNORECASE)
        
        # Convert paragraphs
        content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', content, flags=re.IGNORECASE | re.DOTALL)
        
        # Convert lists
        content = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', content, flags=re.IGNORECASE)
        content = re.sub(r'<ul[^>]*>|</ul>', '', content, flags=re.IGNORECASE)
        content = re.sub(r'<ol[^>]*>|</ol>', '', content, flags=re.IGNORECASE)
        
        # Remove remaining HTML tags
        content = re.sub(r'<[^>]+>', '', content)
        
        # Clean up whitespace
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # Add metadata header
        metadata_header = f"""---
title: {metadata.system_name} - {metadata.template_type.value.replace('_', ' ').title()}
classification: {metadata.classification.value}
version: {metadata.version} 
date: {metadata.created_date.strftime('%Y-%m-%d')}
author: {metadata.author}
organization: {metadata.organization}
---

"""
        
        return metadata_header + content, "md"
    
    def _process_json(self,
                     content: str,
                     metadata: ComplianceMetadata,
                     options: ProcessingOptions) -> Tuple[str, str]:
        """
        Process content to JSON format
        
        Args:
            content: Template content
            metadata: Document metadata
            options: Processing options
            
        Returns:
            Tuple of (json_content, file_extension)
        """
        json_data = {
            'metadata': asdict(metadata),
            'content': content,
            'processing_options': options.to_dict(),
            'generated_date': datetime.now(timezone.utc).isoformat(),
            'content_hash': hashlib.sha256(content.encode()).hexdigest()
        }
        
        return json.dumps(json_data, indent=2, default=str), "json"
    
    def generate_digital_signature(self, content: Union[str, bytes]) -> Optional[str]:
        """
        Generate digital signature for content
        
        Args:
            content: Content to sign
            
        Returns:
            Base64 encoded signature or None if not available
        """
        if not CRYPTO_AVAILABLE or not self.signature_key:
            return None
        
        try:
            # Convert to bytes if string
            if isinstance(content, str):
                content_bytes = content.encode('utf-8')
            else:
                content_bytes = content
            
            # Generate signature
            signature = self.signature_key.sign(
                content_bytes,
                padding.PSS(
                    mgf=padding.MGF1(hashes.SHA256()),
                    salt_length=padding.PSS.MAX_LENGTH
                ),
                hashes.SHA256()
            )
            
            return base64.b64encode(signature).decode('utf-8')
        except Exception as e:
            logger.error(f"Digital signature generation failed: {e}")
            return None
    
    def process_template(self,
                        template_content: str,
                        template_type: TemplateType,
                        metadata: ComplianceMetadata,
                        data: Dict[str, Any],
                        options: ProcessingOptions) -> ProcessingResult:
        """
        Process template with specified options
        
        Args:
            template_content: Rendered template content
            template_type: Type of template
            metadata: Document metadata
            data: Template data
            options: Processing options
            
        Returns:
            ProcessingResult object
        """
        start_time = datetime.now()
        errors = []
        warnings = []
        validation_result = None
        
        try:
            # Validate if requested
            if options.validate_output:
                validation_result = self.validator.validate_template(
                    template_content, template_type, metadata, data
                )
                
                if not validation_result.is_valid:
                    errors.extend(validation_result.errors)
                    warnings.extend(validation_result.warnings)
            
            # Get format processor
            processor = self.format_processors.get(options.output_format)
            if not processor:
                return ProcessingResult(
                    success=False,
                    output_path=None,
                    output_format=options.output_format,
                    file_size=0,
                    processing_time=0.0,
                    content_hash="",
                    validation_result=validation_result,
                    errors=[f"Unsupported output format: {options.output_format}"],
                    warnings=warnings,
                    metadata={}
                )
            
            # Process content
            processed_content, file_ext = processor(template_content, metadata, options)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{metadata.system_id}_{template_type.value}_{timestamp}.{file_ext}"
            output_path = self.output_path / filename
            
            # Save processed content
            if isinstance(processed_content, bytes):
                with open(output_path, 'wb') as f:
                    f.write(processed_content)
                content_hash = hashlib.sha256(processed_content).hexdigest()
            else:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(processed_content)
                content_hash = hashlib.sha256(processed_content.encode()).hexdigest()
            
            # Generate digital signature
            signature = None
            if options.include_signatures:
                signature = self.generate_digital_signature(processed_content)
            
            # Create metadata file
            processing_metadata = {
                'template_type': template_type.value,
                'output_format': options.output_format,
                'processing_options': options.to_dict(),
                'document_metadata': asdict(metadata),
                'processing_date': datetime.now(timezone.utc).isoformat(),
                'content_hash': content_hash,
                'digital_signature': signature,
                'validation_result': validation_result.to_dict() if validation_result else None
            }
            
            metadata_file = output_path.with_suffix(f'.{file_ext}.meta.json')
            with open(metadata_file, 'w') as f:
                json.dump(processing_metadata, f, indent=2, default=str)
            
            # Calculate processing time and file size
            processing_time = (datetime.now() - start_time).total_seconds()
            file_size = output_path.stat().st_size
            
            logger.info(f"Successfully processed template to {options.output_format}: {output_path}")
            
            return ProcessingResult(
                success=True,
                output_path=str(output_path),
                output_format=options.output_format,
                file_size=file_size,
                processing_time=processing_time,
                content_hash=content_hash,
                validation_result=validation_result,
                errors=errors,
                warnings=warnings,
                metadata=processing_metadata
            )
            
        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            logger.error(f"Template processing failed: {e}")
            
            return ProcessingResult(
                success=False,
                output_path=None,
                output_format=options.output_format,
                file_size=0,
                processing_time=processing_time,
                content_hash="",
                validation_result=validation_result,
                errors=[f"Processing failed: {str(e)}"] + errors,
                warnings=warnings,
                metadata={}
            )
    
    def batch_process(self,
                     templates: List[Tuple[str, TemplateType, ComplianceMetadata, Dict[str, Any], ProcessingOptions]],
                     max_workers: int = 4) -> List[ProcessingResult]:
        """
        Process multiple templates in parallel
        
        Args:
            templates: List of template processing tuples
            max_workers: Maximum number of worker threads
            
        Returns:
            List of ProcessingResult objects
        """
        results = []
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks
            future_to_template = {}
            for i, (content, template_type, metadata, data, options) in enumerate(templates):
                future = executor.submit(
                    self.process_template, 
                    content, template_type, metadata, data, options
                )
                future_to_template[future] = i
            
            # Collect results
            for i in range(len(templates)):
                results.append(None)
                
            for future in concurrent.futures.as_completed(future_to_template):
                index = future_to_template[future]
                try:
                    result = future.result()
                    results[index] = result
                except Exception as e:
                    logger.error(f"Batch processing error for template {index}: {e}")
                    results[index] = ProcessingResult(
                        success=False,
                        output_path=None,
                        output_format="unknown",
                        file_size=0,
                        processing_time=0.0,
                        content_hash="",
                        validation_result=None,
                        errors=[f"Batch processing failed: {str(e)}"],
                        warnings=[],
                        metadata={}
                    )
        
        return results


if __name__ == "__main__":
    # Example usage
    import tempfile
    
    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "processed"
        
        processor = TemplateProcessor(output_path)
        
        # Sample template content
        template_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>{{metadata.system_name}} - System Security Plan</title>
        </head>
        <body>
            <h1>{{metadata.system_name}}</h1>
            <h2>Classification: {{metadata.classification.value}}</h2>
            <p>System Description: Test system for demonstration</p>
        </body>
        </html>
        """
        
        from .compliance_template_engine import create_sample_metadata, TemplateType
        metadata = create_sample_metadata("Test System")
        data = {"system_info": {"name": "Test", "description": "Test system"}}
        
        # Test different formats
        formats = processor.get_available_formats()
        for fmt in formats:
            options = ProcessingOptions(
                output_format=fmt,
                include_toc=True,
                classification_headers=True,
                watermark_text="SAMPLE"
            )
            
            result = processor.process_template(
                template_content, TemplateType.SSP, metadata, data, options
            )
            
            print(f"Format {fmt}: {'Success' if result.success else 'Failed'}")
            if result.success:
                print(f"  Output: {result.output_path}")
                print(f"  Size: {result.file_size} bytes")
                print(f"  Time: {result.processing_time:.2f}s")
            else:
                print(f"  Errors: {result.errors}")